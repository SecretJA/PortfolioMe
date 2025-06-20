
from flask import Flask, request, jsonify, send_from_directory
import smtplib
from email.mime.text import MimeText
from email.mime.multipart import MimeMultipart
from datetime import datetime
import json
import os
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

# Email configuration (you'll need to set these as environment variables)
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
EMAIL_ADDRESS = "pphamhoanggiabao19092004@gmail.com"
EMAIL_PASSWORD = os.environ.get('EMAIL_PASSWORD', '')  # Set this in Secrets

def create_word_document(data):
    """Create a Word document with the message data"""
    doc = Document()
    
    # Add header
    header = doc.add_heading('Portfolio Contact Message', 0)
    
    # Add timestamp
    doc.add_paragraph(f"Received: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")
    
    # Add sender information
    doc.add_heading('Sender Information', level=1)
    doc.add_paragraph(f"Name: {data['name']}")
    doc.add_paragraph(f"Email: {data['email']}")
    doc.add_paragraph(f"Subject: {data.get('subject', 'No subject')}")
    doc.add_paragraph("")
    
    # Add message content
    doc.add_heading('Message Content', level=1)
    doc.add_paragraph(data['message'])
    
    # Save document
    filename = f"message_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename

def send_email(data):
    """Send email notification"""
    try:
        msg = MimeMultipart()
        msg['From'] = EMAIL_ADDRESS
        msg['To'] = EMAIL_ADDRESS
        msg['Subject'] = f"Portfolio Contact: {data.get('subject', 'New Message')}"
        
        body = f"""
        New message from your portfolio:
        
        Name: {data['name']}
        Email: {data['email']}
        Subject: {data.get('subject', 'No subject')}
        
        Message:
        {data['message']}
        
        Sent at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        """
        
        msg.attach(MimeText(body, 'plain'))
        
        if EMAIL_PASSWORD:
            with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
                server.starttls()
                server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
                server.send_message(msg)
            return True
        else:
            print("Email password not configured")
            return False
    except Exception as e:
        print(f"Error sending email: {e}")
        return False

@app.route('/')
def index():
    return send_from_directory('.', 'index.html')

@app.route('/<path:filename>')
def serve_static(filename):
    return send_from_directory('.', filename)

@app.route('/send-message', methods=['POST'])
def send_message():
    try:
        data = request.get_json()
        
        # Validate required fields
        if not all(key in data for key in ['name', 'email', 'message']):
            return jsonify({'success': False, 'error': 'Missing required fields'})
        
        # Create Word document
        filename = create_word_document(data)
        
        # Send email notification
        email_sent = send_email(data)
        
        # Save to JSON file as backup
        backup_data = {
            'timestamp': datetime.now().isoformat(),
            'data': data,
            'word_file': filename,
            'email_sent': email_sent
        }
        
        backup_filename = f"messages_backup.json"
        messages = []
        if os.path.exists(backup_filename):
            with open(backup_filename, 'r', encoding='utf-8') as f:
                messages = json.load(f)
        
        messages.append(backup_data)
        
        with open(backup_filename, 'w', encoding='utf-8') as f:
            json.dump(messages, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True, 
            'message': 'Message received and saved',
            'email_sent': email_sent,
            'word_file': filename
        })
        
    except Exception as e:
        print(f"Error processing message: {e}")
        return jsonify({'success': False, 'error': str(e)})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
